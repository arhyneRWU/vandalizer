"""Every tree-runnable corpus validator, proven to fail on a planted defect.

(The three that need the release binaries — `check_references.py`,
`check_scans.py`, and `validate_keys2.py` — have no failure-path test here;
they run in the asset job, against the actual documents, and cannot be
exercised from a fixture the tree job can build.)

A green run says nothing about whether a gate can fail. Each validator here was
shown by hand to reject a specific planted defect during review; those proofs
were one-off, so nothing kept them true. This suite makes them permanent: each
test builds a defect, runs the tool as a subprocess, and asserts both the exit
code and the line the tool prints about it.

CRITICAL: every fixture is generated in `tmp_path` at test time and none of it
is ever committed. A checked-in file holding a retired identity, a name-shaped
string, or a broken key would be scanned by the tree gates themselves and fail
the run it exists to test. For the same reason the strings these tests plant
are assembled from parts here, never written as literals: the person-name scan
reads this file too.

The Office fixtures are built the same way, in `tmp_path`, so that the two
identity scanners are exercised on the shapes the corpus actually shipped: a
name in a body paragraph, a name in a signature-block table cell, a name in a
section header, and a name present only in `docProps` metadata. The table-cell
case is the one that found a live bug — see `SIGNATURE_ROLE_LINE` below.

Run: cd backend && uv run --with pytest pytest \\
       ../benchmarks/corpus/CSU-NSF-001/tools/test_validator_failure_paths.py -q
"""
from __future__ import annotations

import ast
import hashlib
import importlib.util
import json
import shutil
import subprocess
import sys
import tarfile
import zipfile
from pathlib import Path

import pytest

TOOLS = Path(__file__).resolve().parent
KEYS = TOOLS.parent
REPO = TOOLS.parents[3]

# Package name per import name, for the fallback interpreter below.
_PACKAGE = {"pypdf": "pypdf", "docx": "python-docx", "fitz": "pymupdf"}


def interpreter(*modules: str) -> list[str]:
    """Command prefix that can run a tool needing *modules*.

    The suite runs inside the backend environment, which carries PyMuPDF and
    python-docx but not pypdf. Where something is missing, fall back to the
    ephemeral environment the workflow itself uses for those tools, asking for
    every module at once — an ephemeral env has none of them either.
    """
    if all(importlib.util.find_spec(name) is not None for name in modules):
        return [sys.executable]
    if shutil.which("uv") is None:
        raise RuntimeError(
            "these tests invoke the validators via `uv run`; install uv or run "
            "from an env with pypdf/python-docx"
        )
    command = ["uv", "run"]
    for name in modules:
        command += ["--with", _PACKAGE[name]]
    return command + ["python"]


def run(command: list[str], cwd: Path | None = None) -> subprocess.CompletedProcess:
    return subprocess.run(command, cwd=cwd, capture_output=True, text=True)


def tool_constant(filename: str, name: str):
    """A module-level constant read out of a tool's source without importing it.

    Reading rather than importing keeps this suite off the tool's dependencies
    (`validate_identity_safety` needs pypdf, absent from the backend env), and
    keeps the retired identities out of this file — the point of taking the
    denylist from the tool instead of copying a name into a test.
    """
    tree = ast.parse((TOOLS / filename).read_text())
    for node in tree.body:
        targets = node.targets if isinstance(node, ast.Assign) else []
        for target in targets:
            if isinstance(target, ast.Name) and target.id == name:
                return ast.literal_eval(node.value)
    raise AssertionError(f"{filename} no longer defines {name}")


# ---------------------------------------------------------------- verify_assets


def make_tarball(path: Path, member_name: str, payload: bytes) -> None:
    staging = path.parent / f"_staging_{path.name}"
    member = staging / member_name
    member.parent.mkdir(parents=True, exist_ok=True)
    member.write_bytes(payload)
    with tarfile.open(path, "w:gz") as archive:
        archive.add(member, arcname=member_name)
    shutil.rmtree(staging)


def asset_fixture(tmp_path: Path) -> tuple[Path, Path, list[str]]:
    """Two tiny tarballs and a manifest whose digests match them."""
    assets = tmp_path / "assets"
    assets.mkdir()
    names = ["fixture-v0-digital.tar.gz", "fixture-v0-scanned.tar.gz"]
    make_tarball(assets / names[0], "pdf/one.txt", b"digital payload\n")
    make_tarball(assets / names[1], "scanned/light/one.txt", b"scanned payload\n")

    manifest = tmp_path / "manifest.json"
    manifest.write_text(json.dumps({
        "release_assets": {
            "tag": "fixture-v0",
            "assets": [
                {"name": name,
                 "sha256": hashlib.sha256((assets / name).read_bytes()).hexdigest()}
                for name in names
            ],
        }
    }))
    return manifest, assets, names


def verify_assets(manifest: Path, assets: Path) -> subprocess.CompletedProcess:
    return run([sys.executable, str(TOOLS / "verify_assets.py"),
                "--manifest", str(manifest), "--assets-dir", str(assets)])


class TestVerifyAssets:
    def test_matching_digests_pass(self, tmp_path):
        manifest, assets, names = asset_fixture(tmp_path)
        result = verify_assets(manifest, assets)
        assert result.returncode == 0, result.stdout
        assert f"{names[0]}: OK" in result.stdout
        assert f"{names[1]}: OK" in result.stdout

    def test_a_single_corrupted_byte_is_a_mismatch(self, tmp_path):
        manifest, assets, names = asset_fixture(tmp_path)
        target = assets / names[0]
        blob = bytearray(target.read_bytes())
        blob[-1] ^= 0xFF  # one flipped byte, well past the gzip header
        target.write_bytes(bytes(blob))

        result = verify_assets(manifest, assets)
        assert result.returncode == 1, result.stdout
        assert f"{names[0]}: MISMATCH (got " in result.stdout

    def test_an_asset_that_did_not_download_is_missing(self, tmp_path):
        manifest, assets, names = asset_fixture(tmp_path)
        (assets / names[1]).unlink()

        result = verify_assets(manifest, assets)
        assert result.returncode == 1, result.stdout
        assert f"MISSING: {names[1]}" in result.stdout

    def test_a_tarball_absent_from_the_manifest_is_unexpected(self, tmp_path):
        manifest, assets, _names = asset_fixture(tmp_path)
        make_tarball(assets / "fixture-v0-strays.tar.gz", "pdf/x.txt", b"unlisted\n")

        result = verify_assets(manifest, assets)
        assert result.returncode == 1, result.stdout
        assert ("UNEXPECTED (not listed in manifest release_assets): "
                "fixture-v0-strays.tar.gz") in result.stdout


# ------------------------------------------------------------- validate_release


def keys_copy(tmp_path: Path) -> Path:
    """The real in-tree keys, copied so a defect can be planted in them."""
    destination = tmp_path / "keys"
    destination.mkdir()
    for name in ("ground_truth.json", "manifest.json", "benchmark_questions.csv"):
        shutil.copy(KEYS / name, destination / name)
    return destination


def edit_questions(keys: Path, mutate) -> None:
    path = keys / "ground_truth.json"
    payload = json.loads(path.read_text())
    mutate(payload["questions"])
    path.write_text(json.dumps(payload, indent=2))


def validate_release(keys: Path) -> subprocess.CompletedProcess:
    return run(interpreter("pypdf", "docx")
               + [str(TOOLS / "validate_release.py"), "--keys", str(keys)],
               cwd=REPO)


def first_with_corroborating(questions: list[dict]) -> dict:
    for question in questions:
        if question.get("corroborating_sources"):
            return question
    raise AssertionError("no question carries corroborating_sources")


class TestValidateReleaseKeys:
    def test_the_shipped_keys_pass_unmodified(self, tmp_path):
        result = validate_release(keys_copy(tmp_path))
        assert result.returncode == 0, result.stdout + result.stderr
        assert "RELEASE VALIDATION: PASS (keys only;" in result.stdout

    def test_a_corroborating_page_past_the_end_of_the_document_fails(self, tmp_path):
        keys = keys_copy(tmp_path)
        planted = {}

        def plant(questions):
            question = first_with_corroborating(questions)
            question["corroborating_sources"][0][1] = 99
            planted["id"] = question["id"]
            planted["file"] = question["corroborating_sources"][0][0]

        edit_questions(keys, plant)
        result = validate_release(keys)
        assert result.returncode == 1, result.stdout
        assert (f"{planted['id']}: corroborating page 99 outside {planted['file']}"
                in result.stdout)

    def test_a_corroborating_source_duplicating_a_canonical_one_fails(self, tmp_path):
        keys = keys_copy(tmp_path)
        planted = {}

        def plant(questions):
            question = first_with_corroborating(questions)
            canonical = question["sources"][0]
            question["corroborating_sources"].append(list(canonical))
            planted["id"] = question["id"]
            planted["file"] = canonical[0]
            planted["page"] = canonical[1]

        edit_questions(keys, plant)
        result = validate_release(keys)
        assert result.returncode == 1, result.stdout
        assert (f"{planted['id']}: corroborating source duplicates sources: "
                f"{planted['file']} p.{planted['page']}") in result.stdout

    def test_an_unanswerable_question_may_not_carry_corroborating_sources(self, tmp_path):
        keys = keys_copy(tmp_path)
        planted = {}

        def plant(questions):
            question = first_with_corroborating(questions)
            question["answerable"] = False
            planted["id"] = question["id"]

        edit_questions(keys, plant)
        result = validate_release(keys)
        assert result.returncode == 1, result.stdout
        assert (f"{planted['id']}: unanswerable question has corroborating_sources"
                in result.stdout)

    def test_a_question_with_the_key_deleted_fails_rather_than_passing_silently(self, tmp_path):
        keys = keys_copy(tmp_path)
        planted = {}

        def plant(questions):
            question = first_with_corroborating(questions)
            del question["corroborating_sources"]
            planted["id"] = question["id"]

        edit_questions(keys, plant)
        result = validate_release(keys)
        assert result.returncode == 1, result.stdout
        assert f"{planted['id']}: missing corroborating_sources" in result.stdout


# ------------------------------------------------- generated Office fixtures
#
# Built in-process with python-docx and with `zipfile` rather than by shelling
# out to `uv run --with python-docx`: this suite already runs in the backend
# environment, which carries python-docx (see `interpreter()`), so the extra
# process would buy nothing. The .docx has to be built by the same library the
# scanners open it with — a hand-rolled OOXML zip that python-docx cannot read
# would prove nothing about `docx_visible_text()`. The .xlsx is hand-rolled,
# because neither scanner opens a workbook: both read it with `zipfile` only.


def write_split(paragraph, text: str) -> None:
    """Write *text* into *paragraph* as one run per word.

    A name typed as one run also lands verbatim in `word/document.xml`, where
    `office_xml()` would find it — the visible-text walk would then never be
    what failed the run. One run per word keeps every multi-word string out of
    the raw XML, so only the paragraph/table walk can see it; the tests assert
    that with `raw_office_xml()`. Word itself splits runs this readily
    (spell-check state, tracked formatting), so this is not a contrived shape.
    """
    head, _, tail = text.partition(" ")
    paragraph.add_run(head)
    for word in tail.split(" ") if tail else []:
        paragraph.add_run(" " + word)


def build_docx(path: Path, paragraphs: tuple[str, ...] = (),
               cell_lines: tuple[str, ...] = (), header_lines: tuple[str, ...] = (),
               author: str | None = None) -> Path:
    """A .docx with body paragraphs, a one-cell table, a header, and an author.

    `cell_lines` become the paragraphs of a single table cell, which is the
    signature-block shape: a name over the role line beneath it. `header_lines`
    go in the first section's header, where nothing in the body walk can see
    them.
    """
    from docx import Document

    document = Document()
    for text in paragraphs:
        write_split(document.add_paragraph(), text)
    if cell_lines:
        cell = document.add_table(rows=1, cols=1).rows[0].cells[0]
        for index, text in enumerate(cell_lines):
            write_split(cell.paragraphs[0] if index == 0 else cell.add_paragraph(), text)
    if header_lines:
        header = document.sections[0].header
        header.is_linked_to_previous = False
        for index, text in enumerate(header_lines):
            write_split(header.paragraphs[0] if index == 0 else header.add_paragraph(),
                        text)
    if author is not None:
        document.core_properties.author = author
    document.save(str(path))
    return path


# Interpolated rather than written out: the content-types attribute name, next
# to the element name in front of it, reads as First-Last to the in-tree
# person-name gate, which scans this file like any other.
_EXTENSION = "Extension"
_XLSX_CONTENT_TYPES = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
    f'<Default {_EXTENSION}="rels" ContentType="application/vnd.openxmlformats-package'
    f'.relationships+xml"/><Default {_EXTENSION}="xml" ContentType="application/xml"/>'
    '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxml'
    'formats-officedocument.spreadsheetml.sheet.main+xml"/>'
    '<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd'
    '.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
    '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxml'
    'formats-package.core-properties+xml"/></Types>'
)
_XLSX_ROOT_RELS = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
    '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument'
    '/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
    '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006'
    '/relationships/metadata/core-properties" Target="docProps/core.xml"/></Relationships>'
)
_XLSX_WORKBOOK = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
    '<sheets><sheet name="Sheet1" sheetId="1" r:id="rId1"/></sheets></workbook>'
)
_XLSX_WORKBOOK_RELS = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
    '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument'
    '/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/></Relationships>'
)
_XLSX_SHEET = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
    '<sheetData/></worksheet>'
)


def build_xlsx(path: Path, creator: str, last_modified_by: str) -> Path:
    """A minimal workbook whose only interesting content is `docProps/core.xml`."""
    core = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006'
        '/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/">'
        f"<dc:creator>{creator}</dc:creator>"
        f"<cp:lastModifiedBy>{last_modified_by}</cp:lastModifiedBy>"
        "</cp:coreProperties>"
    )
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("[Content_Types].xml", _XLSX_CONTENT_TYPES)
        archive.writestr("_rels/.rels", _XLSX_ROOT_RELS)
        archive.writestr("xl/workbook.xml", _XLSX_WORKBOOK)
        archive.writestr("xl/_rels/workbook.xml.rels", _XLSX_WORKBOOK_RELS)
        archive.writestr("xl/worksheets/sheet1.xml", _XLSX_SHEET)
        archive.writestr("docProps/core.xml", core)
    return path


def raw_office_xml(path: Path) -> str:
    """What `office_xml()` sees — used to prove it is *not* what caught a name."""
    with zipfile.ZipFile(path) as archive:
        return "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith((".xml", ".rels"))
        )


# ------------------------------------------------------- validate_identity_safety

LEGACY_NAMES = tool_constant("validate_identity_safety.py", "LEGACY_NAMES")
REQUIRED_IDS = tool_constant("validate_identity_safety.py", "REQUIRED_IDS")


def identity_package(tmp_path: Path, body: str) -> Path:
    package = tmp_path / "package"
    package.mkdir()
    roles = "\n".join(f"- role identifier {role}" for role in REQUIRED_IDS)
    (package / "notes.md").write_text(f"# Fixture\n\n{roles}\n\n{body}\n")
    return package


def validate_identity(package: Path) -> subprocess.CompletedProcess:
    return run(interpreter("pypdf", "docx")
               + [str(TOOLS / "validate_identity_safety.py"), str(package)],
               cwd=REPO)


class TestValidateIdentitySafety:
    def test_a_retired_identity_anywhere_in_the_package_fails(self, tmp_path):
        retired = LEGACY_NAMES[0]
        package = identity_package(tmp_path, f"The award is held by {retired}.")

        result = validate_identity(package)
        assert result.returncode == 1, result.stdout
        assert "IDENTITY SAFETY: FAIL" in result.stdout
        assert f"notes.md: forbidden legacy identity '{retired}'" in result.stdout

    def test_role_identifiers_and_no_retired_identity_pass(self, tmp_path):
        package = identity_package(tmp_path, "The award is held by the role above.")

        result = validate_identity(package)
        assert result.returncode == 0, result.stdout + result.stderr
        assert "IDENTITY SAFETY: PASS" in result.stdout


# The first retired identity that carries a given name and a surname: the
# two-part shape is what `write_split()` needs to keep out of the raw XML.
TWO_PART_LEGACY = next(name for name in LEGACY_NAMES if " " in name)

CLEAN_BODY = "The award is held by the role above."

# Interpolated, not written out: a literal honorific followed by a capitalised
# token is a person-name finding in this file, even when the token is a role id.
VPR_ID = next(role for role in REQUIRED_IDS if "VPR" in role)


class TestValidateIdentitySafetyOfficeFiles:
    """The .docx and .xlsx branches of `scan()`, one planted defect each.

    Word text lives in four places the tool reads differently, and the corpus's
    own v0.3.2 shipped a retired name in them: a biosketch body paragraph, a
    signature-block table cell, a section header, and document metadata. The
    first three are only reachable through `docx_visible_text()`, the fourth
    only through `office_xml()`; each test asserts the other path could not have
    been what failed the run.
    """

    def test_a_retired_identity_in_a_docx_body_paragraph_fails(self, tmp_path):
        package = identity_package(tmp_path, CLEAN_BODY)
        document = build_docx(
            package / "biosketch.docx",
            paragraphs=(f"{TWO_PART_LEGACY} — Associate Professor",
                        "Role identifier CSU-PI-001."),
        )
        # Split across runs, so `office_xml()` cannot see the joined-up name.
        assert TWO_PART_LEGACY not in raw_office_xml(document)

        result = validate_identity(package)
        assert result.returncode == 1, result.stdout + result.stderr
        assert (f"biosketch.docx: forbidden legacy identity '{TWO_PART_LEGACY}'"
                in result.stdout)

    def test_a_retired_identity_in_a_docx_table_cell_fails(self, tmp_path):
        """The signature-block shape: the branch that dies if table walking breaks."""
        package = identity_package(tmp_path, CLEAN_BODY)
        document = build_docx(
            package / "signatures.docx",
            paragraphs=("Institutional approval.",),
            cell_lines=(f"Dr. {TWO_PART_LEGACY}", "Vice President for Research"),
        )
        assert TWO_PART_LEGACY not in raw_office_xml(document)

        result = validate_identity(package)
        assert result.returncode == 1, result.stdout + result.stderr
        assert (f"signatures.docx: forbidden legacy identity '{TWO_PART_LEGACY}'"
                in result.stdout)

    def test_a_retired_identity_in_a_docx_section_header_fails(self, tmp_path):
        """Headers and footers are their own walk, and their own XML part."""
        package = identity_package(tmp_path, CLEAN_BODY)
        document = build_docx(
            package / "cover.docx",
            paragraphs=("Institutional cover sheet.",),
            header_lines=(TWO_PART_LEGACY,),
        )
        assert TWO_PART_LEGACY not in raw_office_xml(document)

        result = validate_identity(package)
        assert result.returncode == 1, result.stdout + result.stderr
        assert (f"cover.docx: forbidden legacy identity '{TWO_PART_LEGACY}'"
                in result.stdout)

    def test_a_retired_identity_only_in_docx_metadata_fails(self, tmp_path):
        """Nothing visible carries the name; `office_xml()` reads docProps."""
        package = identity_package(tmp_path, CLEAN_BODY)
        build_docx(
            package / "authored.docx",
            paragraphs=("Prepared for CSU-PI-001 by the role above.",),
            author=TWO_PART_LEGACY,
        )

        result = validate_identity(package)
        assert result.returncode == 1, result.stdout + result.stderr
        assert (f"authored.docx: forbidden legacy identity '{TWO_PART_LEGACY}'"
                in result.stdout)

    def test_a_docx_carrying_role_identifiers_only_passes(self, tmp_path):
        package = identity_package(tmp_path, CLEAN_BODY)
        build_docx(
            package / "signatures.docx",
            paragraphs=("Institutional approval.",),
            cell_lines=(f"Dr. {VPR_ID}", "Vice President for Research"),
            author="CSU-PI-001",
        )

        result = validate_identity(package)
        assert result.returncode == 0, result.stdout + result.stderr
        assert "IDENTITY SAFETY: PASS" in result.stdout

    def test_a_retired_identity_only_in_xlsx_metadata_fails(self, tmp_path):
        """A workbook has no `docx_visible_text()`; `office_xml()` is the only pass."""
        package = identity_package(tmp_path, CLEAN_BODY)
        build_xlsx(package / "budget.xlsx",
                   creator="CSU-PI-001", last_modified_by=TWO_PART_LEGACY)

        result = validate_identity(package)
        assert result.returncode == 1, result.stdout + result.stderr
        assert (f"budget.xlsx: forbidden legacy identity '{TWO_PART_LEGACY}'"
                in result.stdout)

    def test_an_xlsx_carrying_role_identifiers_only_passes(self, tmp_path):
        package = identity_package(tmp_path, CLEAN_BODY)
        build_xlsx(package / "budget.xlsx",
                   creator="CSU-PI-001", last_modified_by="CSU-COI-001")

        result = validate_identity(package)
        assert result.returncode == 0, result.stdout + result.stderr
        assert "IDENTITY SAFETY: PASS" in result.stdout


# ------------------------------------------------------------ scan_person_names

# Assembled from parts, never written as one string: this file is itself scanned
# by the in-tree person-name gate, and a name-shaped literal here would be a new
# finding in it. Neither token appears in the tool's NON_PERSON_TOKENS, so the
# pair is exactly what the scanner is built to surface.
NON_PERSON_TOKENS = tool_constant("scan_person_names.py", "NON_PERSON_TOKENS")
_INVENTED = ("Torvald", "Quimbly")
# Also invented, nothing to do with LEGACY_NAMES: it stands for a baseline entry
# whose finding has since left the corpus, which must warn rather than fail.
_STALE_BASELINE_NAME = ("Marla", "Fennwick")
INVENTED_NAME = " ".join(_INVENTED)
STALE_NAME = " ".join(_STALE_BASELINE_NAME)


@pytest.mark.parametrize("token", _INVENTED + _STALE_BASELINE_NAME)
def test_the_planted_tokens_are_not_excused_by_the_tool(token):
    assert token not in NON_PERSON_TOKENS


def scan_package(tmp_path: Path) -> Path:
    package = tmp_path / "scan"
    package.mkdir()
    (package / "notes.md").write_text(
        f"# Fixture\n\n{INVENTED_NAME} reviewed the sampling plan.\n"
    )
    return package


def write_baseline(tmp_path: Path, entries: list[list[str]]) -> Path:
    path = tmp_path / "baseline.json"  # outside the scanned package, deliberately
    path.write_text(json.dumps(entries, indent=2))
    return path


def scan_names(package: Path, baseline: Path | None = None) -> subprocess.CompletedProcess:
    command = interpreter("fitz", "docx") + [str(TOOLS / "scan_person_names.py"),
                                             str(package)]
    if baseline is not None:
        command += ["--baseline", str(baseline)]
    return run(command, cwd=REPO / "backend")


class TestScanPersonNames:
    def test_without_a_baseline_any_candidate_fails(self, tmp_path):
        result = scan_names(scan_package(tmp_path))
        assert result.returncode == 1, result.stdout + result.stderr
        assert "PERSON-NAME SCAN: 1 candidate(s) to review" in result.stdout
        assert f"- notes.md [text] {INVENTED_NAME!r}" in result.stdout

    def test_a_baseline_holding_the_finding_passes(self, tmp_path):
        package = scan_package(tmp_path)
        baseline = write_baseline(tmp_path, [
            ["notes.md", "text", INVENTED_NAME, "context reviewed by a human"],
        ])
        result = scan_names(package, baseline)
        assert result.returncode == 0, result.stdout + result.stderr
        assert "PERSON-NAME SCAN: PASS — every finding is a reviewed baseline entry" \
            in result.stdout

    def test_a_stale_baseline_entry_warns_but_does_not_fail(self, tmp_path):
        package = scan_package(tmp_path)
        baseline = write_baseline(tmp_path, [
            ["notes.md", "text", INVENTED_NAME, "context reviewed by a human"],
            ["gone.md", "text", STALE_NAME, "a name that left the corpus"],
        ])
        result = scan_names(package, baseline)
        assert result.returncode == 0, result.stdout + result.stderr
        assert (f"warning: baseline entry no longer found — gone.md [text] "
                f"{STALE_NAME!r}") in result.stdout

    def test_a_baseline_missing_the_finding_fails(self, tmp_path):
        package = scan_package(tmp_path)
        baseline = write_baseline(tmp_path, [
            ["gone.md", "text", STALE_NAME, "a name that left the corpus"],
        ])
        result = scan_names(package, baseline)
        assert result.returncode == 1, result.stdout + result.stderr
        assert "PERSON-NAME SCAN: FAIL — 1 finding(s) not in the baseline" in result.stdout
        assert f"- notes.md [text] {INVENTED_NAME!r}" in result.stdout


# A lone surname behind an honorific. `NAME` needs two adjacent capitalised
# words of three letters or more, so it cannot match "Quimbly, chair …"; the
# `HONORIFIC` pass is the only thing in `candidates()` that can, which is what
# makes this the case that pins it. The comma bounds the reported tail.
HONORIFIC_SURNAME = _INVENTED[1]
HONORIFIC_HIT = f"Dr. {HONORIFIC_SURNAME}"
HONORIFIC_LINE = f"{HONORIFIC_HIT}, chair of the review panel."

ROLE_LINE = "Role identifier CSU-PI-001."

# The true signature-block shape, and the regression this pins: `docx_lines()`
# yields a table cell as one string, newlines and all, and with `\s+` between
# the groups of `NAME` this matched First-Middle-Last *across* the line break,
# landed the last group on "Vice", and dropped the candidate as a non-person
# token — the name went unreported. `NAME` now separates on `[ \t]+`, so the
# match cannot leave the line the name is on.
SIGNATURE_ROLE_LINE = "Vice President for Research"


def office_package(tmp_path: Path) -> Path:
    """An empty package for exactly one generated Office file, so counts are exact."""
    package = tmp_path / "scan_office"
    package.mkdir()
    return package


class TestScanPersonNamesOfficeFiles:
    """The .docx and .xlsx branches of `collect()`.

    `docx_lines()` walks paragraphs, then tables, then headers and footers;
    `office_metadata()` reads `docProps/*`. Each fixture puts its one
    name-shaped string in a single one of those places, so the finding can only
    have come from that walk.
    """

    def test_a_name_in_a_docx_table_cell_is_a_finding(self, tmp_path):
        package = office_package(tmp_path)
        build_docx(package / "signatures.docx",
                   paragraphs=(ROLE_LINE,),
                   cell_lines=(INVENTED_NAME, SIGNATURE_ROLE_LINE))

        result = scan_names(package)
        assert result.returncode == 1, result.stdout + result.stderr
        assert "PERSON-NAME SCAN: 1 candidate(s) to review" in result.stdout
        assert f"- signatures.docx [body] {INVENTED_NAME!r}" in result.stdout

    def test_a_baseline_holding_the_table_cell_finding_passes(self, tmp_path):
        package = office_package(tmp_path)
        build_docx(package / "signatures.docx",
                   paragraphs=(ROLE_LINE,),
                   cell_lines=(INVENTED_NAME, SIGNATURE_ROLE_LINE))
        baseline = write_baseline(tmp_path, [
            ["signatures.docx", "body", INVENTED_NAME, "context reviewed by a human"],
        ])

        result = scan_names(package, baseline)
        assert result.returncode == 0, result.stdout + result.stderr
        assert "PERSON-NAME SCAN: PASS — every finding is a reviewed baseline entry" \
            in result.stdout

    def test_a_name_in_a_docx_section_header_is_a_finding(self, tmp_path):
        """`docx_lines()` walks headers and footers after paragraphs and tables."""
        package = office_package(tmp_path)
        build_docx(package / "cover.docx",
                   paragraphs=(ROLE_LINE,), header_lines=(INVENTED_NAME,))

        result = scan_names(package)
        assert result.returncode == 1, result.stdout + result.stderr
        assert "PERSON-NAME SCAN: 1 candidate(s) to review" in result.stdout
        assert f"- cover.docx [body] {INVENTED_NAME!r}" in result.stdout

    def test_an_honorific_flags_a_surname_the_name_regex_cannot_match(self, tmp_path):
        package = office_package(tmp_path)
        build_docx(package / "signatures.docx",
                   paragraphs=(ROLE_LINE, HONORIFIC_LINE))

        result = scan_names(package)
        assert result.returncode == 1, result.stdout + result.stderr
        assert "PERSON-NAME SCAN: 1 candidate(s) to review" in result.stdout
        assert f"- signatures.docx [body] {HONORIFIC_HIT!r}" in result.stdout

    def test_a_name_only_in_docx_metadata_is_tagged_metadata(self, tmp_path):
        package = office_package(tmp_path)
        build_docx(package / "authored.docx",
                   paragraphs=(ROLE_LINE,), author=INVENTED_NAME)

        result = scan_names(package)
        assert result.returncode == 1, result.stdout + result.stderr
        assert "PERSON-NAME SCAN: 1 candidate(s) to review" in result.stdout
        assert f"- authored.docx [metadata] {INVENTED_NAME!r}" in result.stdout
        assert f"docProps/core.xml:dc:creator={INVENTED_NAME}" in result.stdout

    def test_a_name_in_xlsx_metadata_is_tagged_metadata(self, tmp_path):
        """A workbook contributes no lines at all — metadata is its whole surface."""
        package = office_package(tmp_path)
        build_xlsx(package / "budget.xlsx",
                   creator="CSU-PI-001", last_modified_by=INVENTED_NAME)

        result = scan_names(package)
        assert result.returncode == 1, result.stdout + result.stderr
        assert "PERSON-NAME SCAN: 1 candidate(s) to review" in result.stdout
        assert f"- budget.xlsx [metadata] {INVENTED_NAME!r}" in result.stdout


# ---------------------------------------------------------------- validate_keys

FIGURE = "$1,234,567.89"
FIGURE_DOC = "A_Figure_Doc.pdf"
PROSE_DOC = "B_Prose_Doc.pdf"
# A reasoned answer: no distinctive figure, and no phrase the cited document
# restates, so string matching cannot speak to it either way. That is the shape
# of citation KNOWN_UNVERIFIABLE exists to pin.
REASONED_ANSWER = "No. Service-center charges are excluded under the policy in force."


def write_pdf(path: Path, lines: list[str]) -> None:
    import pymupdf

    document = pymupdf.open()
    page = document.new_page()
    height = 72
    for line in lines:
        page.insert_text((72, height), line)
        height += 18
    document.save(str(path))
    document.close()


def keys_fixture(tmp_path: Path) -> tuple[Path, Path]:
    """One verifiable citation and one that no string match can reach."""
    binaries = tmp_path / "bin"
    (binaries / "pdf").mkdir(parents=True)
    # Created as layout documentation only: the loader globs `source/*.xlsx`,
    # and a missing directory yields no matches rather than raising, so a
    # fixture citing no workbook would work without it.
    (binaries / "source").mkdir()
    write_pdf(binaries / "pdf" / FIGURE_DOC,
              ["Award summary", f"Total amount requested: {FIGURE}"])
    write_pdf(binaries / "pdf" / PROSE_DOC,
              ["Policy narrative", "This page carries narrative prose only."])

    keys = tmp_path / "keys"
    keys.mkdir()
    (keys / "ground_truth.json").write_text(json.dumps({
        "version": "fixture",
        "questions": [
            {"id": "Q901", "answerable": True, "answer": FIGURE,
             "sources": [[FIGURE_DOC, 1, "header line"]]},
            {"id": "Q902", "answerable": True, "answer": REASONED_ANSWER,
             "sources": [[PROSE_DOC, 1, "narrative"]]},
        ],
    }))
    return keys, binaries


def validate_keys(keys: Path, binaries: Path, *extra: str) -> subprocess.CompletedProcess:
    return run([sys.executable, str(TOOLS / "validate_keys.py"),
                "--keys", str(keys), "--binaries", str(binaries), *extra],
               cwd=REPO / "backend")


class TestValidateKeysUnverifiablePin:
    def test_an_unpinned_unverifiable_citation_fails_the_run(self, tmp_path):
        keys, binaries = keys_fixture(tmp_path)
        result = validate_keys(keys, binaries)
        assert result.returncode == 1, result.stdout + result.stderr
        # The verifiable half is checked on its cited page, so the failure is
        # only about the citation no pass can reach.
        assert "citations verified on the cited page: 1" in result.stdout
        assert f"NEW   Q902 {PROSE_DOC} p.1" in result.stdout
        assert "KEY VALIDATION: FAIL" in result.stdout
        assert ("1 citation(s) checked by neither this pass nor validate_keys2.py, "
                "and not pinned in KNOWN_UNVERIFIABLE") in result.stdout

    def test_the_same_citation_passes_once_it_is_named(self, tmp_path):
        keys, binaries = keys_fixture(tmp_path)
        result = validate_keys(keys, binaries,
                               "--allow-unverifiable", f"Q902:{PROSE_DOC}")
        assert result.returncode == 0, result.stdout + result.stderr
        assert f"known Q902 {PROSE_DOC} p.1" in result.stdout
        assert "KEY VALIDATION: PASS" in result.stdout
