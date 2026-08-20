#!/usr/bin/env python3
"""Release gate for synthetic identities in CSU-NSF-001.

A denylist of the natural-person names v0.3.2 used, plus the role identifiers
that replaced them. It proves the retired names are gone; it cannot catch a
*new* real-person name, which is what `scan_person_names.py` is for.

Takes one or more directories so the same gate covers both halves of the
release: the keys as they sit in the tree, and the unpacked documents.

Run: uv run --with pypdf --with python-docx python \\
       benchmarks/corpus/CSU-NSF-001/tools/validate_identity_safety.py \\
       benchmarks/corpus/CSU-NSF-001 [DIR ...]
"""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


# The natural-person names earlier corpus versions used for synthetic roles,
# retired in v0.3.3; the denylist has to spell them out in order to test for them.
LEGACY_NAMES = [
    "Elena Maris", "David Okafor", "Priya Nolen", "Samuel Reed", "Andrew Rhyne",
    "Sedgemoor", "Obuya", "Vasquez-Osei", "Bellweather", "Okonkwo-Hale", "Ryland",
    "Draymore", "Nakagawa-Pruitt", "Anstruther", "Corliss", "Tanaka-Boyd", "Pemberly",
]
REQUIRED_IDS = ["CSU-PI-001", "CSU-COI-001", "CSU-VPR-001", "FED-NEG-001"]
FORBIDDEN_FAKE_CITATION_MARKERS = [
    "(fictional)",
    "All products listed are fictional works",
]


def docx_visible_text(path: Path) -> str:
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        chunks.extend(p.text for p in section.header.paragraphs)
        chunks.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(chunks)


def office_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith((".xml", ".rels"))
        )


def pdf_text(path: Path) -> str:
    return "\n".join((page.extract_text() or "") for page in PdfReader(path).pages)


def scan(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return docx_visible_text(path) + "\n" + office_xml(path)
    if path.suffix.lower() == ".xlsx":
        return office_xml(path)
    if path.suffix.lower() == ".pdf":
        return pdf_text(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("packages", type=Path, nargs="+",
                        help="one or more directories to scan (keys dir, unpacked assets)")
    args = parser.parse_args()

    failures: list[str] = []
    corpus_text = []
    scanned = 0

    for package in args.packages:
        files = [
            path for path in package.rglob("*")
            if path.is_file()
            and path.suffix.lower() in {".docx", ".pdf", ".xlsx", ".json", ".csv", ".md"}
        ]
        scanned += len(files)
        for path in sorted(files):
            text = scan(path)
            corpus_text.append(text)
            for name in LEGACY_NAMES:
                if name in text:
                    failures.append(
                        f"{path.relative_to(package)}: forbidden legacy identity '{name}'"
                    )
            if path.name == "10_Biographical_Sketches.docx":
                for marker in FORBIDDEN_FAKE_CITATION_MARKERS:
                    if marker in text:
                        failures.append(
                            f"{path.name}: legacy fictional-citation marker '{marker}'"
                        )
                product_ids = re.findall(r"SYN-PUB-(?:PI|COI)-\d{3}", docx_visible_text(path))
                if len(set(product_ids)) != 20:
                    failures.append(
                        f"{path.name}: expected 20 unique SYN-PUB records, "
                        f"found {len(set(product_ids))}"
                    )

    all_text = "\n".join(corpus_text)
    for role_id in REQUIRED_IDS:
        if role_id not in all_text:
            failures.append(f"required role identifier missing: {role_id}")
    if failures:
        print("IDENTITY SAFETY: FAIL")
        print("\n".join(f"- {failure}" for failure in failures))
        sys.exit(1)
    print(f"IDENTITY SAFETY: PASS ({scanned} files scanned)")


if __name__ == "__main__":
    main()
